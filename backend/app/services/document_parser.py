from io import BytesIO

import pymupdf
from docx import Document


class DocumentExtractionError(Exception):
    """Raised when text cannot be extracted from a document."""


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from every page of a PDF stored in memory."""

    if not file_bytes.startswith(b"%PDF"):
        raise DocumentExtractionError("The uploaded file is not a valid PDF.")

    try:
        document = pymupdf.open(stream=file_bytes, filetype="pdf")

        pages = []

        for page in document:
            page_text = page.get_text("text").strip()

            if page_text:
                pages.append(page_text)

        document.close()

        return "\n\n".join(pages).strip()

    except Exception as error:
        raise DocumentExtractionError(
            "The PDF could not be read."
        ) from error


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract paragraph and table text from a DOCX file."""

    try:
        document = Document(BytesIO(file_bytes))
        content = []

        # Extract normal paragraphs.
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                content.append(text)

        # Many resumes use tables for layout, so extract table cells too.
        for table in document.tables:
            for row in table.rows:
                row_values = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text and cell_text not in row_values:
                        row_values.append(cell_text)

                if row_values:
                    content.append(" | ".join(row_values))

        return "\n".join(content).strip()

    except Exception as error:
        raise DocumentExtractionError(
            "The DOCX file could not be read."
        ) from error


def extract_document_text(
    filename: str,
    file_bytes: bytes,
) -> tuple[str, str]:
    """Select the correct extraction method from the filename."""

    lowered_filename = filename.lower()

    if lowered_filename.endswith(".pdf"):
        return extract_pdf_text(file_bytes), "pdf"

    if lowered_filename.endswith(".docx"):
        return extract_docx_text(file_bytes), "docx"

    raise DocumentExtractionError(
        "Only PDF and DOCX resumes are supported."
    )